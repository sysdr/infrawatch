import os
from datetime import datetime, timedelta
from typing import Dict, List
from sqlalchemy.orm import Session
from models.kpi import KPIMetric, ExecutiveReport
from services.kpi_calculator import KPICalculator
from services.trend_analyzer import TrendAnalyzer

class ReportGenerator:
    def __init__(self, db: Session):
        self.db = db
        self.kpi_calculator = KPICalculator(db)
        self.trend_analyzer = TrendAnalyzer(db)

    def generate_executive_report(self, report_type: str = "weekly", output_dir: str = "reports") -> Dict:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import RGBColor
        end_date = datetime.utcnow()
        start_date = end_date - timedelta(days=7 if report_type == "weekly" else (1 if report_type == "daily" else 30))
        doc = Document()
        doc.add_heading('Executive Summary Report', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph()
        p.add_run(f"Period: {start_date.strftime('%Y-%m-%d')} to {end_date.strftime('%Y-%m-%d')}").bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        kpis = self.kpi_calculator.get_dashboard_kpis()
        doc.add_heading('Key Performance Indicators', 1)
        key_insights = []
        for kpi in kpis[:5]:
            kpi_para = doc.add_paragraph()
            kpi_para.add_run(f"{kpi['display_name']}: ").bold = True
            value_text = f"{kpi['current_value']:.2f}" if kpi.get('current_value') is not None else "N/A"
            if kpi.get('unit'):
                value_text += f" {kpi['unit']}"
            kpi_para.add_run(value_text)
            if kpi.get('change_percentage') is not None:
                change_run = kpi_para.add_run(f" ({kpi['change_percentage']:+.1f}%)")
                change_run.font.color.rgb = RGBColor(0, 128, 0) if kpi['change_percentage'] > 0 else RGBColor(255, 0, 0)
            try:
                trend = self.trend_analyzer.analyze_trend(kpi['metric_name'])
                if trend.get('is_anomaly'):
                    key_insights.append({"metric": kpi['display_name'], "insight": f"Anomaly detected: {trend['trend_direction']} trend", "severity": "high"})
            except Exception:
                pass
        if key_insights:
            doc.add_heading('Key Insights', 1)
            for insight in key_insights:
                para = doc.add_paragraph(style='List Bullet')
                para.add_run(f"{insight['metric']}: ").bold = True
                para.add_run(insight['insight'])
        os.makedirs(output_dir, exist_ok=True)
        filepath = os.path.join(output_dir, f"{report_type}_report_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.docx")
        doc.save(filepath)
        self.db.add(ExecutiveReport(
            report_type=report_type, period_start=start_date, period_end=end_date,
            file_path=filepath, summary={"total_kpis": len(kpis), "anomalies_detected": len(key_insights)}, key_insights=key_insights
        ))
        self.db.commit()
        return {"report_type": report_type, "filepath": filepath, "period_start": start_date.isoformat(), "period_end": end_date.isoformat(), "kpis_analyzed": len(kpis), "insights_count": len(key_insights)}
